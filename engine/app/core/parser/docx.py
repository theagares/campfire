"""
app/core/parser/docx.py
DOCX 파싱 — python-docx. 파이프라인/server/pipeline/parsers/docx.py 이식.
"""

from __future__ import annotations

import io


def extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)
