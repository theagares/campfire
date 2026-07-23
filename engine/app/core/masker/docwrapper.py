"""
app/core/masker/docwrapper.py
마스킹된 텍스트를 파일로 래핑한다.
파이프라인/extension/utils/docwrapper.js 의 이식:
  - DOCX 래핑(한국어 Malgun Gothic), PDF 입력도 _masked.docx 로 출력 (PLAN §6).
  - MD(text/plain) 래핑도 제공(파서가 이미 레이아웃을 버리므로 AI 이해에 유리).

JS 원본은 손수 OOXML ZIP 을 만들었지만, 파이썬 엔진은 의존성으로 이미 python-docx
를 갖고 있으므로 그걸로 안정적으로 생성한다(동일 결과: 문단별 Malgun Gothic).
"""

from __future__ import annotations

import io
import re

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _masked_name(orig_file_name: str, new_ext: str) -> str:
    """원본 파일명 → *_masked.<new_ext>. (docwrapper.js 파일명 규칙 이식)"""
    if "." in orig_file_name:
        base = orig_file_name.rsplit(".", 1)[0]
    else:
        base = orig_file_name or "masked_document"
    new_ext = new_ext.lstrip(".")
    return f"{base}_masked.{new_ext}"


def wrap_as_docx(masked_text: str, orig_file_name: str) -> dict:
    """마스킹 텍스트를 DOCX 바이트로 래핑. 반환 {bytes, mime_type, file_name}."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    for line in masked_text.split("\n"):
        para = doc.add_paragraph()
        if line.strip():
            run = para.add_run(line)
            run.font.name = "Malgun Gothic"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:eastAsia", "w:hAnsi"):
                rfonts.set(qn(attr), "Malgun Gothic")

    buf = io.BytesIO()
    doc.save(buf)
    return {
        "bytes": buf.getvalue(),
        "mime_type": DOCX_MIME,
        "file_name": _masked_name(orig_file_name, "docx"),
    }


def wrap_as_md(masked_text: str, orig_file_name: str) -> dict:
    """마스킹 텍스트를 .md(text/plain) 로 래핑 (docwrapper.js wrapAsMd 이식)."""
    return {
        "bytes": masked_text.encode("utf-8"),
        "mime_type": "text/plain",
        "file_name": _masked_name(orig_file_name, "md"),
    }


def wrap_masked_file(masked_text: str, orig_file_name: str, fmt: str = "docx") -> dict:
    """공개 API. PLAN §6 대로 기본은 DOCX(PDF 입력도 _masked.docx)."""
    if fmt == "md":
        return wrap_as_md(masked_text, orig_file_name)
    return wrap_as_docx(masked_text, orig_file_name)
